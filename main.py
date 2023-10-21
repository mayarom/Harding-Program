# Importing necessary modules
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import Flask, request, render_template, send_file, after_this_request, jsonify
import os
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import logging

# Creating a new Flask web application
app = Flask(__name__, template_folder='templates')


# Function to get the current date in a specific format
def get_current_date():
    return datetime.now().strftime('%Y-%m-%d')  # Returns the current date as a string (YYYY-MM-DD)


# Function to generate a file path for output files based on the current date
def get_output_file_path(date):
    return os.path.join('uploads', f'Harding - {date}.docx')  # Constructs a file path in the 'uploads' directory


# Function to extract non-empty lines from a DOCX file
def get_lines_from_docx(docx_path):
    try:
        doc = Document(docx_path)  # Opens a DOCX file
        # Returns a list of non-empty lines from the DOCX file
        return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    except Exception as e:
        logging.exception(f"Failed to load {docx_path}: {str(e)}")  # Logs an exception if the file cannot be opened
        return []


# Function to process the input text file and produce an output DOCX file
def process_txt_file(input_file_path, output_file_path, os_lines):
    document = Document()  # Creates a new DOCX document

    with open(input_file_path, 'r') as file:  # Opens the text file in read mode
        for line in file:  # Iterates over each line in the text file
            clean_line = line.strip().replace("ERROR", "")  # Removes the word "ERROR" from the line

            if "ERROR" in line:  # Checks if the line contains the word "ERROR"
                para = document.add_paragraph(clean_line)  # Adds the line to the DOCX document
                is_line_in_os = any(clean_line.strip() in os_line.strip() for os_line in os_lines)

                if is_line_in_os:  # If the line matches any OS line
                    for run in para.runs:
                        run.bold = True  # Sets the text to bold

                    for os_line in os_lines:
                        if clean_line.strip() in os_line:
                            # Decides which document to add based on the content of os_line
                            source_docx = 'Win12.docx'  # Default
                            if "windows16" in os_line:
                                source_docx = 'Win16.docx'
                            elif "windows19" in os_line:
                                source_docx = 'Win19.docx'

                            # Adds content to the document
                            add_page_to_document(os_line, document, os.path.join('docs', source_docx))

                    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Aligns the paragraph to the right
                else:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Sets the text color to red if no match is found

    document.save(output_file_path)  # Saves the DOCX document
    logging.info(f"File saved as {output_file_path}")  # Logs the save operation


# Function to add a page to the document with specific formatting
def add_page_to_document(page_content, document, source_docx):
    source = Document(source_docx)  # Opens the source DOCX file

    for paragraph in source.paragraphs:
        if page_content.strip() in paragraph.text:
            new_paragraph = document.add_paragraph(paragraph.text)  # Adds the paragraph to the document

            new_run = new_paragraph.runs[0]
            new_run.bold = True  # Sets the text to bold
            new_run.font.size = Pt(12)  # Sets the font size to 12
            new_run.font.color.rgb = RGBColor(0, 0, 255)  # Sets the text color to blue

            run_element = new_run._r
            rpr = run_element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn('w:eastAsia'), 'Arial')  # Sets the font style

            new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Aligns the paragraph to the right


# Route for the home page, which accepts both GET and POST requests
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':  # If the server received a POST request
        input_file = request.files['input_file']  # Gets the uploaded file
        os_choice = request.form.get('os_choice')  # Retrieves the OS choice from the form
        logging.info(f"Operating system chosen: {os_choice}")

        if input_file and os_choice:  # If both a file and OS choice were received
            input_file_path = os.path.join('uploads', 'input.txt')
            input_file.save(input_file_path)  # Saves the uploaded file

            os_docx_mapping = {  # Dictionary mapping OS choices to document filenames
                "windows12": "Win12.docx",
                "windows16": "Win16.docx",
                "windows19": "Win19.docx"
            }

            docx_filename = os_docx_mapping.get(os_choice)  # Gets the document filename based on the OS choice
            if docx_filename:  # If a valid OS choice was made
                docx_path = os.path.join('docs', docx_filename)  # Constructs the path to the DOCX file
                os_lines = get_lines_from_docx(docx_path)  # Retrieves lines from the DOCX file

                current_date = get_current_date()  # Gets the current date
                output_file_path = get_output_file_path(current_date)  # Constructs the output file path

                # Processes the text file and creates the output DOCX file
                process_txt_file(input_file_path, output_file_path, os_lines)

                # Decorator to execute the function after the current request
                @after_this_request
                def delete_files(response):
                    try:
                        os.remove(input_file_path)  # Deletes the input file after processing
                        os.remove(output_file_path)  # Deletes the output file after sending
                        logging.info("Temporary files deleted successfully")  # Logs the deletion
                    except Exception as e:
                        # Logs any exceptions encountered during deletion
                        logging.exception("Error removing or closing downloaded file handle: ", exc_info=e)
                    return response  # Returns the response object

                return send_file(output_file_path, as_attachment=True)  # Sends the output file as a download

            else:
                logging.warning(f"Invalid OS choice: {os_choice}")  # Logs an invalid OS choice
                return "Invalid OS choice", 400  # Returns an error response

    return render_template('home.html')  # Renders the home page template for GET requests


# Checks if the script is executed directly and not imported as a module in another script
if __name__ == '__main__':
    if not os.path.exists('uploads'):  # If the 'uploads' directory does not exist
        os.makedirs('uploads')  # Creates the 'uploads' directory
    # Configures the logging
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    # Runs the Flask application
    app.run(host='0.0.0.0', port=5000, debug=True)


# Route for processing data, accepts only POST requests
@app.route('/process', methods=['POST'])
def process():
    response_data = "It's working!"  # Response message
    return jsonify({'message': response_data})  # Returns a JSON response
